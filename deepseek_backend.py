import os
import re
import sys
import time
from datetime import date
import requests
import urllib3
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from google import genai

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
time.sleep(2)


def build_prompt(property_description: str) -> str: 
     return ( 
         f"你是一位高成交導向的房仲銷售顧問。請根據以下物件資訊，精準產出三個渠道的銷售文案與一個規格解析。\n\n" 
         "【輸出格式要求】\n" 
         "你必須嚴格依序輸出以下四個標籤區塊，不可遺漏，且標籤文字必須完全一致：\n\n" 
         "物件規格解析\n" 
         "(請在此處條列物件的精準地段、建築規劃、公設比、在售格局、定價與車位、管理費與完工日期等核心數據。若資料缺失請標示「未提供」)\n\n" 
         "【591 專業版】\n" 
         "你必須嚴格按照以下六段式骨架分段輸出，漏掉任何一段將視為嚴重的程式碼錯誤：\n" 
         "**[1. 一句抓眼球的標題]**：要包含地標或最大優勢。\n" 
         "**[2. 一句最強差異]**：直接點出這間房跟同區其他物件不一樣的特點。\n" 
         "**[3. 三個買方在意的理由]**：針對物件優勢條列。\n" 
         "**[4. 一段生活場景]**：用描寫式文字寫出居住感。\n" 
         "**[5. 一段完整規格]**：列出坪數、格局、樓層、管理費、車位類型，缺一不可，不足標示「未提供」。\n" 
         "**[6. 一個明確邀請]**：結尾強制寫「本週僅釋出 3 組帶看名額，歡迎私訊看照片並預約賞屋時段」。\n\n" 
         "【FB 社團吸粉版】\n" 
         "(請將上述六段式核心轉譯為適合 Facebook 轉發、具備高社群感染力、口語化且精簡的爆款文案，同樣必須包含明確邀請結尾。)\n\n" 
         "【LINE/限動秒殺版】\n" 
         "(請將核心資訊極度精簡壓縮，適合 LINE 群組與限時動態秒殺快速閱讀，同樣必須包含明確邀請結尾。)\n\n" 
         "【執行規則】\n" 
         "1. 文案要口語化、精簡，絕對不要出現「格局方正、採光佳」這種無效廢話。\n" 
         "2. 【硬性禁止】絕對禁止輸出任何內部筆記、狀態提示、檢查備註（例如：數據需人工確認、等待核對、文案生成中）。\n" 
         "3. 【禁輸出後台評級】絕對禁止輸出任何系統排名、網站人氣榜、熱門度標籤（例如：嚴禁出現『南港人氣榜第3名』或類似排名文字）。\n" 
         "4. 【抗性畫面化】嚴禁直接輸出冰冷的抗性數字（如：公設比 37.5%），必須將其轉譯為買方嚮往的畫面。範例：『純住 37 戶精緻隱私城堡，高公設比換來的是高得房質感的梯廳與純粹的鄰里環境，把干擾留給外面，把尊榮留給自己。』\n" 
         "5. 【指令鎖死】文案結尾的明確邀請，你必須 100% 複製並輸出這句話，不准更改任何一個字：『本週僅釋出 3 組帶看名額，歡迎私訊看照片並預約賞屋時段』。\n" 
         "6. 【直接產出】你的輸出必須直接作為最終交付給客戶的內容，不要包含任何自我解釋、道歉或備註。\n\n" 
         f"待處理資訊：{property_description}" 
     )


def resolve_input_text(user_input: str) -> tuple[str, str]:
    if "591.com.tw" not in user_input:
        return user_input, ""
    print("正在解析網址...")
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML,繼 Gecko) Chrome/125.0 Safari/537.36",
    }
    response = requests.get(user_input, headers=headers, timeout=15, verify=False)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, "html.parser")
    title = soup.title.get_text(strip=True) if soup.title else ""
    text = soup.get_text(separator=" ", strip=True)
    return text[:2000], title


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]", "_", value)
    cleaned = re.sub(r"\s+", "_", cleaned).strip("_")
    return cleaned or "物件"


def split_inputs(raw_input: str) -> list[str]:
    parts = re.split(r"[,\n]+", raw_input)
    return [item.strip() for item in parts if item.strip()]


def extract_sections(content: str) -> dict[str, str]: 
     """超強容錯正則表達式，完全相容有無方括號、多餘空格或字串變體""" 
     sections = { 
         "【591 專業版】": "", 
         "【FB 社團吸粉版】": "", 
         "【LINE/限動秒殺版】": "", 
     } 
     
     # 模糊匹配各渠道區塊，並切分內容 
     p591 = re.search(r"(591\s*專業描述優化版|591\s*專業版|591\s*精簡版|【591\s*專業版】|591)(.*?)(?=◆FB|【FB|FB\s*社團|◆LINE|【LINE|LINE|◆限動|$)", content, re.DOTALL | re.IGNORECASE) 
     pfb = re.search(r"(FB\s*社團爆款版|FB\s*社團|FB\s*爆款|【FB\s*社團吸粉版】|FB)(.*?)(?=◆LINE|【LINE|LINE|◆限動|限動|$)", content, re.DOTALL | re.IGNORECASE) 
     pline = re.search(r"(LINE\s*/\s*限動秒殺版|LINE\s*限動|LINE|限動|【LINE/限動秒殺版】)(.*?)$", content, re.DOTALL | re.IGNORECASE) 
     
     if p591: sections["【591 專業版】"] = p591.group(2).strip() 
     if pfb: sections["【FB 社團吸粉版】"] = pfb.group(2).strip() 
     if pline: sections["【LINE/限動秒殺版】"] = pline.group(2).strip() 
     
     return sections 


def filter_generic_landmarks(lines: list[str]) -> list[str]:
    blocked = ["便利商店", "加油站"]
    filtered: list[str] = []
    for line in lines:
        if any(keyword in line for keyword in blocked):
            continue
        filtered.append(line)
    return filtered


def limit_lines(lines: list[str], max_lines: int) -> list[str]:
    return lines[:max_lines]


def extract_intel_section(content: str) -> str: 
     """模糊匹配物件規格解析區塊，找不到就拿整篇大文本的前 1000 字當作保險""" 
     markers = ["物件規格解析", "【物件核心規格與黃金價值】", "物件規格", "規格解析", "核心規格"] 
     for marker in markers: 
         start_index = content.find(marker) 
         if start_index != -1: 
             # 截取到下一個大區塊之前 
             end_index = content.find("591", start_index) 
             if end_index == -1: 
                 end_index = content.find("【591", start_index) 
             if end_index != -1: 
                 return content[start_index:end_index] 
             return content[start_index:] 
     return content[:1000] 


def clean_markdown(text: str) -> str:
    return re.sub(r"\*+", "", text)


def add_field(paragraph, instruction: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    paragraph._p.append(field)


def add_paragraph_with_highlight(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pattern = re.compile(r"(總價[^，。\n]*|單價[^，。\n]*|完工日期[^，。\n]*)")
    remaining = text
    last_index = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_index:
            paragraph.add_run(text[last_index:start])
        highlighted = paragraph.add_run(text[start:end])
        highlighted.font.highlight_color = WD_COLOR_INDEX.YELLOW
        last_index = end
    if last_index < len(text):
        paragraph.add_run(text[last_index:])


def apply_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    heading1 = document.styles["Heading 1"]
    heading1.font.name = "Microsoft JhengHei"
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    heading1._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    heading2 = document.styles["Heading 2"]
    heading2.font.name = "Microsoft JhengHei"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0, 0, 0)
    heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    heading2._element.rPr.rFonts.set(qn("w:ascii"), "Arial")


def set_header_footer(document: Document, property_name: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = True
    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)
    left_cell.text = f"[{property_name}]"
    right_cell.text = date.today().strftime("%Y/%m/%d")
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.add_run("第 ")
    add_field(footer_paragraph, "PAGE")
    footer_paragraph.add_run(" 頁 / 共 ")
    add_field(footer_paragraph, "NUMPAGES")
    footer_paragraph.add_run(" 頁")


def extract_region_and_name(title: str, fallback: str) -> tuple[str, str]:
    region_candidates = [
        "台北市", "新北市", "桃園市", "台中市", "台南市", "高雄市", "基隆市", "新竹市", "嘉義市",
        "台北", "新北", "桃園", "台中", "台南", "高雄", "基隆", "新竹", "嘉義",
        "宜蘭", "花蓮", "台東", "雲林", "彰化", "南投", "屏東", "苗栗",
    ]
    region = "全台"
    for candidate in region_candidates:
        if candidate in title:
            region = candidate
            break
    name_source = title or fallback
    name_source = re.sub(r"591|租|售|買|房屋|物件|出售", "", name_source)
    name_source = re.sub(r"[\[\]\(\)【】]", "", name_source)
    name = sanitize_filename(name_source)[:12] or "物件"
    return region, name


def generate_listing(client: genai.Client, property_description: str) -> str:
    prompt = build_prompt(property_description)
    model_sequence = [
        "models/gemini-2.5-flash",
        "models/gemini-2.0-flash",
        "models/gemini-1.5-flash",
        "models/gemma-3-12b",
    ]
    last_error: Exception | None = None
    for model_name in model_sequence:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=prompt,
                config={"temperature": 0.7},
            )
            content = getattr(response, "text", None)
            if content:
                return content.strip()
            last_error = ValueError("Gemini response content is empty")
        except Exception as error:
            last_error = error
            continue
    if last_error:
        raise last_error
    raise ValueError("Gemini response content is empty")


def save_docx(title: str, content: str) -> str: 
     # 1. 徹底清洗與過濾內部提示 
     content = remove_internal_prompts(content) 
     content = re.sub(r"\n\s*\n+", "\n\n", content) 
     cleaned_content = clean_markdown(content) 
     cleaned_content = remove_internal_prompts(cleaned_content) 
     
     # 2. 提取區塊 
     intel_section = clean_markdown(extract_intel_section(cleaned_content)) 
     intel_section = remove_internal_prompts(intel_section) 
     sections = extract_sections(cleaned_content) 
     
     output_dir = os.path.join(os.getcwd(), "Outputs") 
     os.makedirs(output_dir, exist_ok=True) 
     
     document = Document() 
     apply_styles(document) 
     region, property_name = extract_region_and_name(title, intel_section[:12]) 
     set_header_footer(document, property_name) 
 
     # 3. 物件規格解析頁面 
     document.add_heading("物件規格解析", level=1) 
     document.add_paragraph("#物件規格 #市場解析 #銷售要點") 
     table = document.add_table(rows=2, cols=2) 
     table.cell(0, 0).text = "本案單價" 
     table.cell(0, 1).text = "區域行情" 
     table.cell(1, 0).text = "價差判讀" 
     table.cell(1, 1).text = "詳見下方說明" 
     
     # 終極降級防線：如果截取太短，說明切片失敗，直接拿整篇當備份，拒絕預留字！ 
     display_intel = intel_section if len(intel_section.strip()) > 30 else cleaned_content 
     intel_lines = limit_lines(filter_generic_landmarks(display_intel.splitlines()), 15) 
     for line in intel_lines: 
         line_strip = line.strip() 
         if line_strip and not any(h in line_strip for h in ["物件規格解析", "【物件核心規格與黃金價值】", "物件規格", "規格解析"]): 
             add_paragraph_with_highlight(document, line_strip) 
     document.add_page_break() 
 
     # 4. 591 專業版（防禦：抓不到直接用整篇 cleaned_content 代替，絕不留白、絕不印出生成中） 
     document.add_heading("591 專業版", level=1) 
     document.add_paragraph("#成交戰術 #數據精準 #專業建議") 
     professional_content = sections["【591 專業版】"].strip() if sections["【591 專業版】"].strip() else cleaned_content 
     professional_lines = limit_lines(filter_generic_landmarks(professional_content.splitlines()), 16) 
     for line in professional_lines: 
         line_strip = line.strip() 
         if line_strip: 
             add_paragraph_with_highlight(document, line_strip) 
     document.add_page_break() 
 
     # 5. FB 社團吸粉版（防禦：抓不到直接用整篇 cleaned_content 代替） 
     document.add_heading("FB 社團吸粉版", level=1) 
     document.add_paragraph("#在地社群 #吸粉曝光 #熱區生活") 
     fb_content = sections["【FB 社團吸粉版】"].strip() if sections["【FB 社團吸粉版】"].strip() else cleaned_content 
     fb_lines = limit_lines(filter_generic_landmarks(fb_content.splitlines()), 12) 
     for line in fb_lines: 
         line_strip = line.strip() 
         if line_strip: 
             add_paragraph_with_highlight(document, line_strip) 
     document.add_page_break() 
 
     # 6. LINE / 限動版（防禦：抓不到直接用整篇 cleaned_content 代替） 
     document.add_heading("LINE/限動秒殺版", level=1) 
     document.add_paragraph("#VIP急售 #限量釋出 #稀缺搶手") 
     line_content = sections["【LINE/限動秒殺版】"].strip() if sections["【LINE/限動秒殺版】"].strip() else cleaned_content 
     line_lines = limit_lines(filter_generic_landmarks(line_content.splitlines()), 8) 
     for line in line_lines: 
         line_strip = line.strip() 
         if line_strip: 
             add_paragraph_with_highlight(document, line_strip) 
 
     # 7. 聯絡區塊 
     contact_para = document.add_paragraph() 
     contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER 
     contact_name, contact_phone = get_contact_info() 
     contact_text = f"業務聯繫人：{contact_name} | 電話：{contact_phone}" 
     contact_run = contact_para.add_run(contact_text) 
     contact_run.bold = True 
 
     filename = f"【成品報告】{region}_{property_name}_{date.today().strftime('%m%d')}.docx" 
     file_path = os.path.join(output_dir, filename) 
     document.save(file_path) 
     return file_path 



def remove_internal_prompts(text: str) -> str:
    """徹底清洗與過濾內部提示與標籤垃圾"""
    if not text:
        return ""
    # 移除硬性禁止的內部筆記與提示
    patterns = [
        r"數據需人工確認",
        r"等待核對",
        r"文案生成中",
        r"內部筆記",
        r"檢查備註",
        r"系統排名",
        r"網站人氣榜",
        r"熱門度標籤",
    ]
    for p in patterns:
        text = re.sub(p, "", text)
    # 移除星號與井號 (Markdown 殘留)
    text = text.replace("*", "").replace("#", "")
    return text.strip()


def get_contact_info() -> tuple[str, str]:
    """獲取預設聯繫資訊 (供 DOCX 版本使用)"""
    return "您的稱呼", "您的聯絡電話"


def main() -> None:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise SystemExit("請設定 GEMINI_API_KEY")
    client = genai.Client(api_key=api_key)
    user_input = input("請直接貼上 591 網址或房屋資訊：")
    inputs = split_inputs(user_input)
    if not inputs:
        raise SystemExit("請提供房地產物件描述文字")
    for item in inputs:
        description, title = resolve_input_text(item)
        if not description:
            continue
        output = generate_listing(client, description)
        base_title = title or (description[:12] if description else "物件")
        save_docx(base_title, output)
        print("✅ 超級專家報告已生成在 Outputs 資料夾")


if __name__ == "__main__":
    main()
